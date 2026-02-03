import streamlit as st
import os
import base64
import pandas as pd
from io import BytesIO
from fpdf import FPDF
from docx import Document
from streamlit_gsheets import GSheetsConnection

# 1. Page Config
st.set_page_config(page_title="Exclusive Holidays SL", page_icon="✈️", layout="wide")

# --- DATABASE CONNECTION ---
conn = st.connection("gsheets", type=GSheetsConnection)

def load_user_db():
    try:
        df = conn.read(worksheet="Sheet1", ttl=0)
        return df
    except:
        return None

def update_user_password(username, new_password):
    df = conn.read(worksheet="Sheet1", ttl=0)
    df.loc[df['username'] == username, ['password', 'status']] = [str(new_password), "Active"]
    conn.update(worksheet="Sheet1", data=df)
    st.cache_data.clear()

def add_user_to_db(new_u, new_p):
    df = conn.read(worksheet="Sheet1", ttl=0)
    new_row = pd.DataFrame([{"username": str(new_u), "password": str(new_p), "status": "New"}])
    updated_df = pd.concat([df, new_row], ignore_index=True)
    conn.update(worksheet="Sheet1", data=updated_df)
    st.cache_data.clear()

def remove_user_from_db(username):
    df = conn.read(worksheet="Sheet1", ttl=0)
    updated_df = df[df['username'] != username]
    conn.update(worksheet="Sheet1", data=updated_df)
    st.cache_data.clear()

# Initialize Session States
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
if 'current_user' not in st.session_state:
    st.session_state.current_user = None
if 'itinerary' not in st.session_state:
    st.session_state.itinerary = []
if 'needs_password_change' not in st.session_state:
    st.session_state.needs_password_change = False
if 'admin_form_key' not in st.session_state:
    st.session_state.admin_form_key = 0
if 'builder_form_key' not in st.session_state:
    st.session_state.builder_form_key = 0

def get_base64(bin_file):
    if os.path.exists(bin_file):
        with open(bin_file, 'rb') as f:
            data = f.read()
        return base64.b64encode(data).decode()
    return None

# 2. CSS Styling
bg_img = "https://images.unsplash.com/photo-1586500036706-41963de24d8b?q=80&w=2574&auto=format&fit=crop"
st.markdown(f"""
    <style>
    [data-testid="stAppViewContainer"] {{
        background: linear-gradient(rgba(0, 0, 0, 0.4), rgba(0, 0, 0, 0.4)), url("{bg_img}");
        background-size: cover; background-position: center; background-attachment: fixed;
    }}
    div[data-testid="stVerticalBlock"] > div:has(> div:empty) {{ display: none !important; }}
    [data-testid="stVerticalBlock"] {{ background-color: transparent !important; }}
    .stTextInput input, .stTextArea textarea {{ background-color: rgba(255, 255, 255, 0.95) !important; color: #1e1e1e !important; }}
    ::placeholder {{ color: #555555 !important; opacity: 1; }}
    .stButton > button {{ color: #000000 !important; font-weight: 800 !important; background-color: #ffffff !important; }}
    h1, h2, h3, p, label {{ color: white !important; text-shadow: 2px 2px 4px rgba(0,0,0,0.8); }}
    header {{visibility: hidden;}} footer {{visibility: hidden;}}
    </style>
    """, unsafe_allow_html=True)

def display_branding(logo_size=100, title_size=28, motto_size=14):
    logo_base64 = get_base64("logo.png")
    if logo_base64:
        st.markdown(f'<div style="text-align: center;"><img src="data:image/png;base64,{logo_base64}" width="{logo_size}"></div>', unsafe_allow_html=True)
    st.markdown(f'<h1 style="text-align: center; font-size: {title_size}px; margin-bottom: 0;">EXCLUSIVE HOLIDAYS</h1>', unsafe_allow_html=True)
    st.markdown(f'<p style="text-align: center; font-size: {motto_size}px; font-style: italic; margin-top: 0;">"Unforgettable Island Adventures Awaits"</p>', unsafe_allow_html=True)

# --- PDF LOGIC ---
class ITINERARY_PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 16)
        self.set_text_color(0, 51, 102)
        self.cell(0, 10, 'EXCLUSIVE HOLIDAYS SRI LANKA', 0, 1, 'C')
        self.set_font('Arial', 'I', 10)
        self.cell(0, 10, '"Unforgettable Island Adventures Awaits"', 0, 1, 'C')
        self.ln(10)
        self.line(10, 32, 200, 32)
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def create_pdf(title, itinerary):
    pdf = ITINERARY_PDF()
    pdf.add_page()
    pdf.set_font('Arial', 'B', 14)
    pdf.cell(0, 10, f'Trip Plan: {title}', 0, 1, 'L')
    pdf.ln(5)
    for i, day in enumerate(itinerary):
        pdf.set_fill_color(240, 240, 240)
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 10, f'Day {i+1}: {day["Route"]}', 1, 1, 'L', fill=True)
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 8, f'Distance: {day["Distance"]}  |  Travel Time: {day["Time"]}', 0, 1, 'L')
        pdf.set_font('Arial', '', 11)
        pdf.multi_cell(0, 7, day['Description'])
        pdf.ln(10)
    return pdf.output(dest='S')

# --- WORD LOGIC ---
def to_word(title, itinerary):
    doc = Document()
    doc.add_heading(f'Itinerary: {title}', 0)
    for i, item in enumerate(itinerary):
        doc.add_heading(f'Day {i+1}: {item["Route"]}', level=1)
        doc.add_paragraph(f'Distance: {item["Distance"]} | Time: {item["Time"]}')
        doc.add_paragraph(item['Description'])
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- EXCEL LOGIC ---
def to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.index = df.index + 1
        df.to_excel(writer, sheet_name='Itinerary', index_label='Day')
    return output.getvalue()

# --- LOGIN ---
if not st.session_state.authenticated:
    col1, col2, col3 = st.columns([1, 1.5, 1])
    with col2:
        display_branding(logo_size=200, title_size=42, motto_size=18)
        with st.form("login_form"):
            u_input = st.text_input("Username")
            p_input = st.text_input("Password", type="password")
            if st.form_submit_button("Sign In"):
                df = load_user_db()
                if df is not None:
                    user_row = df[df['username'] == u_input]
                    if not user_row.empty and str(user_row.iloc[0]['password']) == p_input:
                        st.session_state.authenticated = True
                        st.session_state.current_user = u_input
                        if u_input != "admin01" and user_row.iloc[0]['status'] == "New":
                            st.session_state.needs_password_change = True
                        st.rerun()
                    else: st.error("Invalid Credentials")
    st.stop()

# --- FORCED PW CHANGE ---
if st.session_state.needs_password_change:
    display_branding(logo_size=120)
    st.markdown("## 🔒 Set Your Permanent Password")
    with st.form("forced_pw_form"):
        new_p = st.text_input("New Password", type="password")
        conf_p = st.text_input("Confirm New Password", type="password")
        if st.form_submit_button("Update & Continue"):
            if new_p == conf_p and len(new_p) >= 4:
                update_user_password(st.session_state.current_user, new_p)
                st.session_state.needs_password_change = False
                st.rerun()
            else: st.error("Passwords must match (min 4 chars).")
    st.stop()

display_branding(logo_size=80, title_size=24, motto_size=12)

t_col1, t_col2 = st.columns([9, 1])
with t_col2:
    if st.button("Logout"):
        st.session_state.authenticated = False
        st.rerun()

# --- ADMIN PANEL ---
if st.session_state.current_user == "admin01":
    tab1, tab2, tab3 = st.tabs(["Add Staff", "Remove Staff", "Change Admin Password"])
    with tab1:
        st.markdown("### 👨‍💼 Add New Staff")
        new_u = st.text_input("New Username", key=f"user_{st.session_state.admin_form_key}")
        new_p = st.text_input("Temp Password", type="password", key=f"pass_{st.session_state.admin_form_key}")
        if st.button("Register Account"):
            if new_u and new_p:
                add_user_to_db(new_u, new_p)
                st.session_state.admin_form_key += 1
                st.success(f"Added {new_u}!")
                st.rerun()
    with tab2:
        df_users = load_user_db()
        if df_users is not None:
            staff_list = df_users[df_users['username'] != "admin01"]
            for index, row in staff_list.iterrows():
                c_u, c_b = st.columns([3, 1])
                c_u.write(f"👤 **{row['username']}**")
                if c_b.button(f"Delete", key=f"del_{row['username']}"):
                    remove_user_from_db(row['username'])
                    st.rerun()
    with tab3:
        st.markdown("### 🔑 Change Admin Password")
        with st.form("admin_pw_change"):
            old_p = st.text_input("Old Password", type="password")
            new_p = st.text_input("New Password", type="password")
            conf_p = st.text_input("Confirm New Password", type="password")
            if st.form_submit_button("Update Admin Password"):
                df = load_user_db()
                curr_p = str(df[df['username'] == "admin01"].iloc[0]['password'])
                if old_p == curr_p and new_p == conf_p and len(new_p) >= 4:
                    update_user_password("admin01", new_p)
                    st.success("Admin password updated!")
                else: st.error("Verification failed.")

# --- STAFF BUILDER ---
else:
    tab_build, tab_settings = st.tabs(["✈️ Itinerary Builder", "Settings ⚙️"])
    with tab_build:
        tour_title = st.text_input("Tour Title / Client Name", placeholder="e.g. John Smith", key="tour_title_fix")
        c1, c2, c3 = st.columns([2, 1, 1])
        with c1: r_in = st.text_input("Route", key=f"r_{st.session_state.builder_form_key}")
        with c2: d_in = st.text_input("Distance", key=f"d_{st.session_state.builder_form_key}")
        with c3: t_in = st.text_input("Time", key=f"t_{st.session_state.builder_form_key}")
        desc_in = st.text_area("Description", key=f"de_{st.session_state.builder_form_key}")
        
        if st.button("➕ Add Day"):
            if r_in:
                st.session_state.itinerary.append({"Route": r_in, "Distance": d_in, "Time": t_in, "Description": desc_in})
                st.session_state.builder_form_key += 1
                st.rerun()
        if st.button("🗑️ Clear All"):
            st.session_state.itinerary = []
            st.rerun()

        if st.session_state.itinerary:
            st.markdown("---")
            # --- 3 EXPORT BUTTONS ---
            ex1, ex2, ex3 = st.columns(3)
            with ex1:
                st.download_button("📥 Excel", data=to_excel(pd.DataFrame(st.session_state.itinerary)), file_name=f"{tour_title}.xlsx")
            with ex2:
                st.download_button("📥 Word", data=to_word(tour_title, st.session_state.itinerary), file_name=f"{tour_title}.docx")
            with ex3:
                pdf_data = create_pdf(tour_title, st.session_state.itinerary)
                st.download_button("📥 Professional PDF", data=pdf_data, file_name=f"{tour_title}.pdf", mime="application/pdf")

            for i, item in enumerate(st.session_state.itinerary):
                st.markdown(f"**Day {i+1}: {item['Route']}**")
                if st.button(f"Remove Day {i+1}", key=f"rem_{i}"):
                    st.session_state.itinerary.pop(i)
                    st.rerun()
    with tab_settings:
        st.markdown("### 🔑 Change Your Password")
        with st.form("staff_pw_change"):
            old_p = st.text_input("Old Password", type="password")
            new_p = st.text_input("New Password", type="password")
            conf_p = st.text_input("Confirm New Password", type="password")
            if st.form_submit_button("Save New Password"):
                df = load_user_db()
                curr_p = str(df[df['username'] == st.session_state.current_user].iloc[0]['password'])
                if old_p == curr_p and new_p == conf_p and len(new_p) >= 4:
                    update_user_password(st.session_state.current_user, new_p)
                    st.success("Password changed!")
                else: st.error("Verification failed.")
